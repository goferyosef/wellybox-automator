#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""WellyBox Downloader v4.0"""

import re
import threading
import time
import traceback
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    from selenium.webdriver.chrome.service import Service as ChromeService
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.common.exceptions import (
        NoSuchElementException, StaleElementReferenceException,
        TimeoutException,
    )
    from webdriver_manager.chrome import ChromeDriverManager
    SELENIUM_OK = True
except ImportError:
    SELENIUM_OK = False

try:
    import openpyxl
    from openpyxl.styles import PatternFill, Font
    from openpyxl.utils import get_column_letter
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

try:
    import keyring
    KEYRING_OK = True
except ImportError:
    KEYRING_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── Constants ─────────────────────────────────────────────────────────────────
APP_NAME     = "WellyBox Downloader"
KEYRING_SVC  = "WellyBoxApp"
WELLYBOX_URL = "https://app.wellybox.com/"
WAIT_S, WAIT_L = 10, 30

HEBREW_MONTHS = {
    "ינואר":1,"ינו":1,"פברואר":2,"פבר":2,"מרץ":3,"מרס":3,
    "אפריל":4,"אפר":4,"מאי":5,"יוני":6,"יונ":6,
    "יולי":7,"יול":7,"אוגוסט":8,"אוג":8,
    "ספטמבר":9,"ספט":9,"אוקטובר":10,"אוק":10,
    "נובמבר":11,"נוב":11,"דצמבר":12,"דצמ":12,
}
ENGLISH_MONTHS = {
    "january":1,"jan":1,"february":2,"feb":2,"march":3,"mar":3,
    "april":4,"apr":4,"may":5,"june":6,"jun":6,
    "july":7,"jul":7,"august":8,"aug":8,"september":9,"sep":9,
    "october":10,"oct":10,"november":11,"nov":11,"december":12,"dec":12,
}
INVOICE_TYPES = ["חשבונית מס / קבלה", "חשבונית מס/קבלה", "חשבונית"]
RECEIPT_TYPES = ["קבלה"]

# All label strings that appear in the detail panel — never treat these as values
KNOWN_LABELS  = {
    "סוג מסמך", "ספק", "מטבע", "סכום", "תאריך", "מספר", "סטטוס",
    "מקור", "קטגוריה", "הערות", "עסק",
    'תאריך ע"ג המסמך',   # ASCII double-quote variant
    'תאריך ע״ג המסמך',   # Hebrew gershayim variant (U+05F4)
    "תאריך העלאת המסמך", "תאריך הנפקה", "תאריך קבלה",
    "סך הכל", 'סה"כ', 'סה״כ', "מספר חשבונית", "שם העסק",
    "מטבע חשבונית", "מספר מסמך", "פרטים", "קובץ",
    "סטטוס תשלום", "אמצעי תשלום", "סכום לתשלום", "יתרה",
    "תשלום", "הערה", "כתובת", "ח.פ", "ע.מ", "מיקוד",
    "חבר צוות", "מקור", "שולח האימייל",
}

COLOR_GREEN    = "C8E6C9"
COLOR_RED      = "FFCDD2"
COLOR_BLUE     = "BBDEFB"
COLOR_YELLOW   = "FFF9C4"
COLOR_GREY     = "EEEEEE"
COLOR_ORANGE   = "FFE0B2"
COLOR_LAVENDER = "E8EAF6"   # skipped because WellyBox status = saved

CONFIG_PATH = Path(__file__).parent / "wellybox_config.json"

def load_folder_config() -> dict:
    try:
        if CONFIG_PATH.exists():
            import json as _j
            return _j.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

def save_folder_config(cfg: dict) -> None:
    import json as _j
    CONFIG_PATH.write_text(_j.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


# ── Data ──────────────────────────────────────────────────────────────────────
@dataclass
class CardResult:
    idx:      int
    vendor:   str = ""
    doc_date: str = ""
    doc_type: str = ""
    status:   str = ""   # downloaded_invoice|dup_invoice|downloaded_receipt|dup_receipt|skipped|error
    filename: str = ""
    note:     str = ""


# ── Credentials ───────────────────────────────────────────────────────────────
def load_creds():
    if not KEYRING_OK:
        return None, None
    return (keyring.get_password(KEYRING_SVC, "username"),
            keyring.get_password(KEYRING_SVC, "password"))

def save_creds(u, p):
    if KEYRING_OK:
        keyring.set_password(KEYRING_SVC, "username", u)
        keyring.set_password(KEYRING_SVC, "password", p)


# ── Credentials dialog ────────────────────────────────────────────────────────
class CredsDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("פרטי כניסה — WellyBox")
        self.resizable(False, False)
        self.grab_set()
        self.username = ""
        self.password = ""
        saved_u, saved_p = load_creds()
        pad = {"padx": 10, "pady": 5}
        tk.Label(self, text="אימייל:").grid(row=0, column=0, sticky="e", **pad)
        self._u = tk.Entry(self, width=34)
        self._u.grid(row=0, column=1, **pad)
        if saved_u:
            self._u.insert(0, saved_u)
        tk.Label(self, text="סיסמה:").grid(row=1, column=0, sticky="e", **pad)
        self._p = tk.Entry(self, width=34, show="*")
        self._p.grid(row=1, column=1, **pad)
        if saved_p:
            self._p.insert(0, saved_p)
        f = tk.Frame(self)
        f.grid(row=2, column=0, columnspan=2, pady=8)
        tk.Button(f, text="שמור",  width=10, command=self._ok).pack(side="left", padx=6)
        tk.Button(f, text="ביטול", width=10, command=self.destroy).pack(side="left", padx=6)
        self._u.focus_set()
        self.bind("<Return>", lambda _: self._ok())
        self.wait_window()

    def _ok(self):
        self.username = self._u.get().strip()
        self.password = self._p.get()
        self.destroy()


# ── Folder selection dialog ───────────────────────────────────────────────────
class FolderDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("בחר תיקיות שמירה")
        self.resizable(False, False)
        self.grab_set()
        self.invoice_folder: str = ""
        self.receipt_folder: str = ""
        self._confirmed = False

        cfg = load_folder_config()
        default_inv = cfg.get("default_invoice_folder", str(Path.home() / "Desktop"))
        default_rec = cfg.get("default_receipt_folder", str(Path.home() / "Desktop"))

        pad = {"padx": 10, "pady": 6}

        tk.Label(self, text="תיקיית חשבוניות / דוחות:", anchor="e").grid(
            row=0, column=0, sticky="e", **pad)
        self._inv = tk.StringVar(value=default_inv)
        tk.Entry(self, textvariable=self._inv, width=44).grid(row=0, column=1, **pad)
        tk.Button(self, text="…", width=3,
                  command=lambda: self._browse(self._inv)).grid(row=0, column=2, padx=(0, 8))

        self._save_inv = tk.BooleanVar(value=False)
        tk.Checkbutton(self, text="קבע כברירת מחדל לחשבוניות",
                       variable=self._save_inv).grid(
            row=1, column=1, sticky="w", padx=10, pady=(0, 4))

        tk.Label(self, text="תיקיית קבלות:", anchor="e").grid(
            row=2, column=0, sticky="e", **pad)
        self._rec = tk.StringVar(value=default_rec)
        tk.Entry(self, textvariable=self._rec, width=44).grid(row=2, column=1, **pad)
        tk.Button(self, text="…", width=3,
                  command=lambda: self._browse(self._rec)).grid(row=2, column=2, padx=(0, 8))

        self._save_rec = tk.BooleanVar(value=False)
        tk.Checkbutton(self, text="קבע כברירת מחדל לקבלות",
                       variable=self._save_rec).grid(
            row=3, column=1, sticky="w", padx=10, pady=(0, 8))

        bf = tk.Frame(self)
        bf.grid(row=4, column=0, columnspan=3, pady=8)
        tk.Button(bf, text="המשך", width=12,
                  bg="#2e7d32", fg="white", activebackground="#1b5e20",
                  font=("Segoe UI", 10, "bold"),
                  command=self._ok).pack(side="left", padx=6)
        tk.Button(bf, text="ביטול", width=10, command=self.destroy).pack(side="left", padx=6)

        self.bind("<Return>", lambda _: self._ok())
        self.wait_window()

    def _browse(self, var: tk.StringVar):
        current = var.get()
        initial = current if Path(current).exists() else str(Path.home())
        d = filedialog.askdirectory(title="בחר תיקייה", initialdir=initial)
        if d:
            var.set(d)

    def _ok(self):
        inv = self._inv.get().strip()
        rec = self._rec.get().strip()
        if not inv or not rec:
            messagebox.showwarning("קלט חסר", "נא לבחור שתי תיקיות", parent=self)
            return
        self.invoice_folder = inv
        self.receipt_folder = rec
        cfg = load_folder_config()
        if self._save_inv.get():
            cfg["default_invoice_folder"] = inv
        if self._save_rec.get():
            cfg["default_receipt_folder"] = rec
        if self._save_inv.get() or self._save_rec.get():
            save_folder_config(cfg)
        self._confirmed = True
        self.destroy()


# ── Date helpers ──────────────────────────────────────────────────────────────
def parse_date(text: str) -> Optional[datetime]:
    text = text.strip()
    # dd/mm/yyyy or dd.mm.yyyy  (4-digit year — try first to avoid partial match)
    m = re.search(r'(\d{1,2})[/.](\d{1,2})[/.](\d{4})', text)
    if m:
        try:
            return datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except Exception:
            pass
    # dd/mm/yy or dd.mm.yy  (2-digit year)
    m = re.search(r'(\d{1,2})[/.](\d{1,2})[/.](\d{2})(?!\d)', text)
    if m:
        try:
            y = int(m.group(3))
            y = 2000 + y if y < 70 else 1900 + y
            return datetime(y, int(m.group(2)), int(m.group(1)))
        except Exception:
            pass
    # Hebrew: 24 מרץ 2026
    m = re.search(r'(\d{1,2})\s+([א-ת]+)\s+(\d{4})', text)
    if m:
        d, mon, y = int(m.group(1)), m.group(2), int(m.group(3))
        mn = HEBREW_MONTHS.get(mon) or HEBREW_MONTHS.get(mon[:3])
        if mn:
            try:
                return datetime(y, mn, d)
            except Exception:
                pass
    # English: Mar 24th 2026 / Mar 24 2026
    m = re.search(r'([A-Za-z]+)\s+(\d{1,2})(?:st|nd|rd|th)?\s*,?\s*(\d{4})', text)
    if m:
        mon, d, y = m.group(1).lower(), int(m.group(2)), int(m.group(3))
        mn = ENGLISH_MONTHS.get(mon) or ENGLISH_MONTHS.get(mon[:3])
        if mn:
            try:
                return datetime(y, mn, d)
            except Exception:
                pass
    # English: 24 Mar 2026
    m = re.search(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', text)
    if m:
        d, mon, y = int(m.group(1)), m.group(2).lower(), int(m.group(3))
        mn = ENGLISH_MONTHS.get(mon) or ENGLISH_MONTHS.get(mon[:3])
        if mn:
            try:
                return datetime(y, mn, d)
            except Exception:
                pass
    return None

def fmt_date_il(dt: datetime) -> str:
    return f"{dt.day}.{dt.month}.{dt.year}"

def safe_name(s: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', '', s).strip()

def _join_note(existing: str, addition: str) -> str:
    """Append addition to existing note, separated by ' | '."""
    return f"{existing} | {addition}" if existing else addition

def type_matches(doc_type: str, expected_types: list) -> bool:
    """Exact match ignoring spaces and slash variants — prevents 'קבלה' matching 'חשבונית מס / קבלה'."""
    dt = re.sub(r'[\s/]+', '', doc_type).strip()
    for et in expected_types:
        if dt == re.sub(r'[\s/]+', '', et).strip():
            return True
    return False


# ── Bot ───────────────────────────────────────────────────────────────────────
class Bot:
    def __init__(self, days_back: int, invoice_folder: Path, receipt_folder: Path,
                 log_cb, done_cb, max_docs: int = 30, mark_as_saved: bool = False):
        self.days_back      = days_back
        self.invoice_folder = invoice_folder
        self.receipt_folder = receipt_folder
        self._log_cb        = log_cb
        self._done_cb       = done_cb
        self.max_docs       = max_docs
        self.mark_as_saved  = mark_as_saved
        self.driver         = None
        self.stop_event     = threading.Event()
        self.results        = []
        self._lines        = []
        self._cutoff       = (datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
                              - timedelta(days=days_back - 1))

    def start(self):
        threading.Thread(target=self._run, daemon=True).start()

    def stop(self):
        self.stop_event.set()

    def _emit(self, msg: str, level: str = "INFO"):
        ts   = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] [{level}] {msg}"
        self._lines.append(line)
        self._log_cb(line, level)

    # ── main flow ─────────────────────────────────────────────────────────────
    def _run(self):
        try:
            u, p = load_creds()
            if not u or not p:
                self._done_cb(need_creds=True)
                return
            self._emit(f"מתחיל — {self.days_back} ימים (מ-{self._cutoff.strftime('%d/%m/%Y')})")
            self._start_browser()
            self._login(u, p)
            if self.stop_event.is_set():
                return

            # Transfer browser session to requests
            session = self._build_session()

            # Fetch all docs in range via REST API
            self._emit("══════ שולף מסמכים מה-API ══════")
            all_docs = self._fetch_docs(session)
            self._emit(f"נמצאו {len(all_docs)} מסמכים בטווח התאריכים")
            if self.stop_event.is_set():
                return

            # Split by doc_type (API values confirmed from live data)
            invoice_docs = [d for d in all_docs
                            if d.get('doc_type') in ('invoice', 'invoice_and_receipt')]
            receipt_docs = [d for d in all_docs
                            if d.get('doc_type') == 'receipt']

            # Stage 1 — invoices
            self._emit(f"══════ שלב 1: חשבוניות ({len(invoice_docs)}) ══════")
            self._process_docs(invoice_docs, self.invoice_folder, "חשבונית", session)
            if self.stop_event.is_set():
                return

            # Stage 2 — receipts
            self._emit(f"══════ שלב 2: קבלות ({len(receipt_docs)}) ══════")
            self._process_docs(receipt_docs, self.receipt_folder, "קבלה", session)

            self._logout()
            self._save_reports()
            self._emit("✓ סיום בהצלחה")
        except Exception as exc:
            self._emit(f"שגיאה כללית: {exc}", "ERROR")
            self._emit(traceback.format_exc(), "ERROR")
        finally:
            if self.driver:
                try:
                    self.driver.quit()
                except Exception:
                    pass
            self._done_cb(need_creds=False)

    # ── browser ───────────────────────────────────────────────────────────────
    def _start_browser(self):
        self._emit("פותח Chrome…")
        dl_tmp = self.invoice_folder / "_dl_tmp"
        # Clear any leftover files from previous runs so detection works cleanly
        if dl_tmp.exists():
            for f in dl_tmp.iterdir():
                try:
                    f.unlink()
                except Exception:
                    pass
        dl_tmp.mkdir(parents=True, exist_ok=True)
        self._dl_tmp = dl_tmp

        opts = ChromeOptions()
        opts.add_argument("--start-maximized")
        opts.add_argument("--disable-notifications")
        opts.add_argument("--disable-blink-features=AutomationControlled")
        opts.add_experimental_option("excludeSwitches", ["enable-automation"])
        opts.add_experimental_option("useAutomationExtension", False)
        opts.add_experimental_option("prefs", {
            "download.default_directory":  str(dl_tmp),
            "download.prompt_for_download": False,
            "download.directory_upgrade":   True,
            "safebrowsing.enabled":         True,
        })
        service = ChromeService(ChromeDriverManager().install())
        self.driver = webdriver.Chrome(service=service, options=opts)
        self.driver.execute_script(
            "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
        )
        # Set download directory via CDP (works for newer Chrome versions)
        dl_tmp_str = str(dl_tmp).replace("\\", "/")
        try:
            self.driver.execute_cdp_cmd("Browser.setDownloadBehavior", {
                "behavior": "allow",
                "downloadPath": str(dl_tmp),
                "eventsEnabled": True,
            })
        except Exception:
            try:
                self.driver.execute_cdp_cmd("Page.setDownloadBehavior", {
                    "behavior": "allow",
                    "downloadPath": str(dl_tmp),
                })
            except Exception:
                pass
        self._emit(f"  תיקיית הורדות: {dl_tmp}")

    def _w(self, t=WAIT_L):
        return WebDriverWait(self.driver, t)

    def _shot(self, name: str):
        try:
            d = self.invoice_folder / "logs" / "diag"
            d.mkdir(parents=True, exist_ok=True)
            self.driver.save_screenshot(
                str(d / f"{name}_{datetime.now().strftime('%H%M%S')}.png")
            )
        except Exception:
            pass

    # ── login ─────────────────────────────────────────────────────────────────
    def _login(self, username: str, password: str):
        self._emit("מתחבר…")
        self.driver.get(WELLYBOX_URL)
        time.sleep(3)

        # Click כניסה on landing page if present
        try:
            btn = self._w(WAIT_S).until(EC.element_to_be_clickable((
                By.XPATH,
                "//*[self::a or self::button]"
                "[contains(.,'כניסה') or contains(.,'התחברות') or contains(.,'Login')]"
            )))
            btn.click()
            time.sleep(2)
        except TimeoutException:
            pass

        # Email field
        email_el = None
        for xp in ["//input[@type='email']", "//input[@name='email']", "//input[@id='email']"]:
            try:
                email_el = self._w(8).until(EC.element_to_be_clickable((By.XPATH, xp)))
                break
            except TimeoutException:
                continue
        if not email_el:
            raise RuntimeError("שדה אימייל לא נמצא")
        email_el.click()
        time.sleep(0.2)
        email_el.send_keys(Keys.CONTROL + "a")
        email_el.send_keys(username)
        self._emit(f"  אימייל: {username}")

        # Password field (may require submitting email first in 2-step flows)
        pass_el = None
        for attempt in range(2):
            try:
                pass_el = self._w(6).until(
                    EC.element_to_be_clickable((By.XPATH, "//input[@type='password']"))
                )
                break
            except TimeoutException:
                if attempt == 0:
                    try:
                        self.driver.find_element(By.XPATH, "//button[@type='submit']").click()
                        time.sleep(2)
                    except Exception:
                        pass
        if not pass_el:
            raise RuntimeError("שדה סיסמה לא נמצא")
        pass_el.click()
        time.sleep(0.2)
        pass_el.send_keys(Keys.CONTROL + "a")
        pass_el.send_keys(password)

        # Submit
        submitted = False
        for xp in [
            "//button[@type='submit']",
            "//button[contains(.,'התחברות') or contains(.,'כניסה') or contains(.,'Login')]",
        ]:
            try:
                for b in self.driver.find_elements(By.XPATH, xp):
                    if b.is_displayed() and b.is_enabled():
                        self.driver.execute_script("arguments[0].click();", b)
                        submitted = True
                        break
                if submitted:
                    break
            except Exception:
                pass
        if not submitted:
            pass_el.send_keys(Keys.RETURN)

        # Wait for dashboard
        self._emit("ממתין לכניסה…")
        try:
            self._w(WAIT_L).until(EC.any_of(
                EC.url_contains("dashboard"),
                EC.url_contains("receipts"),
                EC.url_contains("ng2ux"),
                EC.presence_of_element_located(
                    (By.XPATH, "//*[contains(.,'כל החשבוניות')]")
                ),
            ))
            self._emit("✓ מחובר")
        except TimeoutException:
            self._shot("login_fail")
            raise RuntimeError("כניסה נכשלה — בדוק פרטי כניסה")
        time.sleep(2)

    # ── API session ───────────────────────────────────────────────────────────
    def _build_session(self):
        import requests, json as _json
        session = requests.Session()

        # Copy cookies without domain restriction so they reach api.app.wellybox.com
        for cookie in self.driver.get_cookies():
            session.cookies.set(cookie['name'], cookie['value'])

        ua = self.driver.execute_script('return navigator.userAgent')
        headers = {
            'User-Agent': ua,
            'Referer':    'https://app.wellybox.com/',
            'Origin':     'https://app.wellybox.com',
        }

        # Try to pull JWT from localStorage / sessionStorage
        token = self.driver.execute_script("""
            var keys = ['access_token','token','auth_token','jwt','id_token',
                        'accessToken','authToken','idToken'];
            for (var s of [localStorage, sessionStorage]) {
                for (var k of keys) {
                    var v = s.getItem(k);
                    if (v) return v;
                }
            }
            return null;
        """)

        if token:
            # Token may itself be a JSON string containing a nested token
            try:
                parsed = _json.loads(token)
                if isinstance(parsed, dict):
                    token = (parsed.get('access_token') or parsed.get('token')
                             or parsed.get('id_token') or token)
            except Exception:
                pass
            headers['Authorization'] = f'Bearer {token}'
            self._emit(f"  ✓ JWT נמצא ({len(token)} תווים)")
        else:
            self._emit("  לא נמצא JWT — מנסה עם cookies בלבד", "WARN")

        session.headers.update(headers)
        self._emit("  ✓ העברתי session ל-requests")
        return session

    # ── fetch all docs via REST API ────────────────────────────────────────────
    def _fetch_docs(self, session) -> list:
        import json
        API_URL = "https://api.app.wellybox.com/api/v1/docs2"
        all_docs = []
        page = 1
        while True:
            if self.stop_event.is_set():
                break
            docs_filter = json.dumps({
                "sort_by": "source_date",
                "sort_dir": "desc",
                "page": page,
                "page_size": 60,
            }, separators=(',', ':'))
            params = {
                "docs_filter": docs_filter,
                "cver": "2.4.1",
                "intent": "screen",
                "ff": "desktop",
                "culture": "he_IL",
            }
            self._emit(f"  שולף עמוד {page}…")
            resp = None
            try:
                resp = session.get(API_URL, params=params, timeout=30)
                resp.raise_for_status()
                data = resp.json()
            except Exception as e:
                self._emit(f"  שגיאת API בעמוד {page}: {e}", "ERROR")
                if resp is not None:
                    self._emit(f"  תשובת שרת: {resp.text[:600]}", "ERROR")
                break

            items = data.get('items') or []
            total_pages = data.get('pages', 1)
            self._emit(f"  עמוד {page}/{total_pages} — {len(items)} פריטים")

            past_cutoff = False
            for item in items:
                if len(all_docs) >= self.max_docs:
                    self._emit(f"  הגעתי למגבלת {self.max_docs} מסמכים, עוצר")
                    past_cutoff = True
                    break
                doc_date_raw = item.get('doc_date') or item.get('source_date') or ''
                doc_dt = None
                if doc_date_raw:
                    try:
                        doc_dt = datetime.fromisoformat(doc_date_raw[:10])
                    except Exception:
                        doc_dt = parse_date(doc_date_raw)

                if doc_dt and doc_dt < self._cutoff:
                    self._emit(f"  תאריך {doc_date_raw[:10]} — מחוץ לטווח, עוצר")
                    past_cutoff = True
                    break
                all_docs.append(item)

            if past_cutoff or page >= total_pages:
                break
            page += 1

        self._emit(f"  סה״כ {len(all_docs)} מסמכים (סרקתי {page} עמודים)")
        return all_docs

    # ── download & save docs ───────────────────────────────────────────────────
    def _process_docs(self, docs: list, dest_folder: Path,
                      stage_name: str, session) -> None:
        dest_folder.mkdir(parents=True, exist_ok=True)
        dl_stat        = f"downloaded_{'invoice' if stage_name == 'חשבונית' else 'receipt'}"
        dup_stat       = f"dup_{'invoice' if stage_name == 'חשבונית' else 'receipt'}"
        collision_stat = f"collision_renamed_{'invoice' if stage_name == 'חשבונית' else 'receipt'}"
        _status_field_logged = False   # log raw status field once per stage for diagnostics

        for idx, doc in enumerate(docs, 1):
            if self.stop_event.is_set():
                break

            vendor = safe_name(doc.get('vendor_title') or 'לא ידוע')

            # Issued date = date printed on the document (תאריך ע"ג המסמך)
            # Never fall back to source_date or doc_date (upload/processing dates) — use "??.??.????" if unknown
            doc_date_raw = (doc.get('issue_date') or doc.get('issued_date')
                            or doc.get('document_date') or doc.get('invoice_date')
                            or doc.get('receipt_date') or '')
            doc_dt = None
            if doc_date_raw:
                try:
                    doc_dt = datetime.fromisoformat(doc_date_raw[:10])
                except Exception:
                    doc_dt = parse_date(doc_date_raw)
            date_str      = fmt_date_il(doc_dt) if doc_dt else '??.??.????'
            date_unknown  = doc_dt is None
            doc_type     = doc.get('doc_type') or ''
            download_url = doc.get('download_pdf_url') or ''
            # Strategy 4: enrich filename with doc number when available
            doc_num_raw  = (doc.get('doc_number') or doc.get('invoice_number')
                            or doc.get('number') or doc.get('doc_num') or '')
            doc_num      = safe_name(str(doc_num_raw)).strip() if doc_num_raw else ''

            # WellyBox card status — try common field names
            wb_status = (doc.get('status') or doc.get('review_status')
                         or doc.get('doc_status') or doc.get('card_status') or '').lower()
            if not _status_field_logged:
                date_fields = {k: doc[k] for k in
                               ('issue_date','issued_date','document_date','invoice_date',
                                'receipt_date','doc_date','source_date','created_at')
                               if k in doc}
                self._emit(f"  [diag] שדות תאריך: {date_fields}")
                status_keys = [k for k in ('status','review_status','doc_status','card_status')
                               if k in doc]
                self._emit(f"  [diag] סטטוס שדה: {status_keys} → '{wb_status}'")
                _status_field_logged = True

            result = CardResult(
                idx=idx,
                vendor=vendor,
                doc_date=date_str,
                doc_type=doc_type,
                note='תאריך לא ברור' if date_unknown else '',
            )

            # Skip "saved" cards; process "new" and unknown statuses
            if wb_status in ('saved', 'נשמר'):
                self._emit(f"  #{idx}: {vendor} {date_str} — סטטוס 'saved', דלג")
                result.status = 'skipped_saved'
                result.note   = 'WellyBox: saved'
                self.results.append(result)
                continue
            if wb_status and wb_status != 'new':
                self._emit(f"  #{idx}: {vendor} {date_str} — סטטוס לא מוכר '{wb_status}', מעבד בכל זאת", "WARN")

            if not download_url:
                self._emit(f"  #{idx}: {vendor} {date_str} — אין קישור הורדה", "WARN")
                result.status = 'skipped'
                result.note   = 'אין קישור הורדה'
                self.results.append(result)
                continue

            base_name = f"{vendor} {date_str}"
            filename  = f"{base_name}.pdf"
            dest_path = dest_folder / filename

            # Download content to memory so we can compare before writing
            self._emit(f"  #{idx}: מוריד {filename}…")
            try:
                resp = session.get(download_url, timeout=60)
                resp.raise_for_status()
                new_content = resp.content
            except Exception as e:
                self._emit(f"  #{idx}: שגיאת הורדה — {e}", "ERROR")
                result.status = 'error'
                result.note   = str(e)
                self.results.append(result)
                continue

            if dest_path.exists():
                # Strategy 3: compare contents — skip only if truly identical
                import hashlib
                existing_hash = hashlib.md5(dest_path.read_bytes()).hexdigest()
                new_hash      = hashlib.md5(new_content).hexdigest()
                if existing_hash == new_hash:
                    self._emit(f"  #{idx}: {filename} — זהה לקיים, דלג")
                    result.status   = dup_stat
                    result.filename = filename
                    result.note     = _join_note(result.note, 'קובץ זהה')
                    self.results.append(result)
                    continue
                # Different content — use doc_num as disambiguator, fall back to counter
                original_name = filename
                if doc_num:
                    filename  = f"{base_name} מס{doc_num}.pdf"
                    dest_path = dest_folder / filename
                counter = 2
                while dest_path.exists():
                    filename  = f"{base_name} ({counter}).pdf"
                    dest_path = dest_folder / filename
                    counter  += 1
                self._emit(f"  #{idx}: תוכן שונה, שומר כ־{filename}")
                result.status = collision_stat
                result.note   = _join_note(result.note, f"שם מקורי: {original_name} — קובץ שונה, הוסף מזהה")
            else:
                result.status = dl_stat

            try:
                dest_path.write_bytes(new_content)
                result.filename = filename
                self._emit(f"  ✓ {filename}")
                if self.mark_as_saved:
                    self._mark_doc_saved(session, doc)
            except Exception as e:
                self._emit(f"  #{idx}: שגיאת כתיבה — {e}", "ERROR")
                result.status = 'error'
                result.note   = _join_note(result.note, str(e))

            self.results.append(result)


    # ── logout ────────────────────────────────────────────────────────────────
    def _logout(self):
        self._emit("מתנתק…")
        try:
            # Try clicking the user/avatar menu then the logout item
            for xp in [
                "//*[contains(@class,'avatar') or contains(@class,'Avatar') or contains(@class,'user-menu')]",
                "//*[@aria-label='account' or @aria-label='user' or @aria-label='profile']",
            ]:
                for el in self.driver.find_elements(By.XPATH, xp):
                    if el.is_displayed():
                        el.click()
                        time.sleep(1)
                        break
            for el in self.driver.find_elements(
                By.XPATH,
                "//*[contains(.,'התנתק') or contains(.,'Logout') or contains(.,'Sign out')]"
            ):
                if el.is_displayed():
                    el.click()
                    self._emit("  ✓ התנתק")
                    return
        except Exception:
            pass

        # Fallback: navigate directly to the logout URL
        try:
            self.driver.get("https://app.wellybox.com/logout")
            time.sleep(2)
            self._emit("  ✓ התנתק (redirect)")
            return
        except Exception:
            pass

        # Browser will be closed immediately after — session ends regardless
        self._emit("  סגירת דפדפן תנתק את החשבון")

    # ── mark doc as saved in WellyBox ─────────────────────────────────────────
    def _mark_doc_saved(self, session, doc):
        doc_id = doc.get('id') or doc.get('_id') or doc.get('doc_id')
        if not doc_id:
            self._emit("  [diag] לא נמצא מזהה מסמך לסימון כנשמר", "WARN")
            return
        try:
            resp = session.patch(
                f"https://api.app.wellybox.com/api/v1/docs/{doc_id}",
                json={"status": "saved"},
                timeout=10,
            )
            if resp.ok:
                self._emit(f"  ✓ סומן כנשמר (id={doc_id})")
            else:
                self._emit(f"  [WARN] סימון כנשמר נכשל: {resp.status_code} {resp.text[:120]}", "WARN")
        except Exception as e:
            self._emit(f"  [WARN] שגיאה בסימון כנשמר: {e}", "WARN")

    # ── reports ───────────────────────────────────────────────────────────────
    def _save_reports(self):
        rep_dir = self.invoice_folder / "דוחות"
        rep_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%d.%m.%Y")

        # ד"ח פעולות — LibreOffice Writer (.docx)
        doc_path = rep_dir / f'ד״ח פעולות {ts}.docx'
        if DOCX_OK:
            doc = Document()
            # Title
            title = doc.add_heading('WellyBox — ד"ח פעולות', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Meta
            doc.add_paragraph(f"תאריך הפעלה: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph(f"ימים אחורה: {self.days_back}")
            doc.add_paragraph(f"תיקיית חשבוניות: {self.invoice_folder}")
            doc.add_paragraph(f"תיקיית קבלות:    {self.receipt_folder}")
            doc.add_paragraph("─" * 60)
            doc.add_paragraph("")
            # Log lines
            doc.add_heading("לוג פעולות", level=2)
            for line in self._lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
                run.font.name = "Courier New"
                if "[ERROR]" in line:
                    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
                elif "[WARN]" in line:
                    run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
            doc.save(str(doc_path))
            self._emit(f'ד״ח פעולות → {doc_path}')
            import os as _os
            try:
                _os.startfile(str(doc_path))
            except Exception:
                pass
        else:
            # Fallback to plain text if python-docx not available
            txt_path = rep_dir / f'ד״ח פעולות {ts}.txt'
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write('WellyBox — ד"ח פעולות\n')
                f.write(f"תאריך הפעלה: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
                f.write(f"ימים אחורה : {self.days_back}\n")
                f.write("=" * 60 + "\n\n")
                for line in self._lines:
                    f.write(line + "\n")
            self._emit(f'ד"ח פעולות → {txt_path}')

        if not OPENPYXL_OK:
            return

        # Excel colored report
        wb  = openpyxl.Workbook()
        ws  = wb.active
        ws.title = 'דו"ח הורדות'

        # Header row
        headers = ["#", "ספק", "תאריך", "סוג מסמך", "סטטוס", "שם קובץ", "הערה"]
        hdr_fill = PatternFill("solid", fgColor="263238")
        hdr_font = Font(bold=True, color="FFFFFF")
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.fill = hdr_fill
            c.font = hdr_font

        STATUS_COLOR = {
            "downloaded_invoice":        COLOR_GREEN,
            "dup_invoice":               COLOR_RED,
            "downloaded_receipt":        COLOR_BLUE,
            "dup_receipt":               COLOR_YELLOW,
            "collision_renamed_invoice": COLOR_YELLOW,
            "collision_renamed_receipt": COLOR_YELLOW,
            "skipped_saved":             COLOR_LAVENDER,
            "skipped":                   COLOR_GREY,
            "error":                     COLOR_ORANGE,
        }
        STATUS_HE = {
            "downloaded_invoice":        "הורד — חשבוניות",
            "dup_invoice":               "כפול — דלג (חשבוניות)",
            "downloaded_receipt":        "הורד — קבלות",
            "dup_receipt":               "כפול — דלג (קבלות)",
            "collision_renamed_invoice": "התנגשות שמות — הורד עם מזהה (חשבוניות)",
            "collision_renamed_receipt": "התנגשות שמות — הורד עם מזהה (קבלות)",
            "skipped_saved":             "דלג — כרטיס מסומן כ-Saved ב-WellyBox",
            "skipped":                   "דלג",
            "error":                     "שגיאה",
        }

        for r in self.results:
            row  = ws.max_row + 1
            fill = PatternFill("solid", fgColor=STATUS_COLOR.get(r.status, COLOR_GREY))
            vals = [r.idx, r.vendor, r.doc_date, r.doc_type,
                    STATUS_HE.get(r.status, r.status), r.filename, r.note]
            for col, v in enumerate(vals, 1):
                c = ws.cell(row=row, column=col, value=v)
                c.fill = fill

        # Legend
        lg = wb.create_sheet("Legend")
        for i, (color, desc) in enumerate([
            (COLOR_GREEN,    "הורד לתיקיית חשבוניות לקליטה"),
            (COLOR_RED,      "קובץ זהה קיים בחשבוניות לקליטה — דלג"),
            (COLOR_BLUE,     "הורד לתיקיית קבלות"),
            (COLOR_YELLOW,   "קובץ זהה קיים בקבלות — דלג"),
            (COLOR_YELLOW,   "התנגשות שמות — קבצים שונים, הורד עם מזהה מסמך (חשבונית/קבלה)"),
            (COLOR_LAVENDER, "דלג — כרטיס מסומן כ-Saved ב-WellyBox (לא חדש)"),
            (COLOR_GREY,     "דלג (סוג/תאריך לא תואם, אין קישור)"),
            (COLOR_ORANGE,   "שגיאה"),
        ], 1):
            lg.cell(row=i, column=1).fill = PatternFill("solid", fgColor=color)
            lg.cell(row=i, column=2, value=desc)

        # Column widths
        for sheet in [ws, lg]:
            for col in sheet.columns:
                w = max((len(str(c.value or "")) for c in col), default=0)
                sheet.column_dimensions[get_column_letter(col[0].column)].width = min(w + 4, 55)

        xls_path = rep_dir / f'דו״ח הורדות {ts}.xlsx'
        wb.save(xls_path)
        self._emit(f'דו"ח הורדות → {xls_path}')


# ── GUI ───────────────────────────────────────────────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.resizable(False, False)
        _icon = Path(__file__).parent / "wellybox_icon.ico"
        if _icon.exists():
            try:
                self.iconbitmap(str(_icon))
            except Exception:
                pass
        self._bot = None
        self._build_ui()
        # Restore persisted mark-as-saved preference
        cfg = load_folder_config()
        if cfg.get('remember_mark_saved', False):
            self._remember_mark.set(True)
            self._mark_saved.set(cfg.get('mark_as_saved', False))
        # Show ready message — do NOT auto-start anything
        self.after(200, self._show_welcome)

    def _show_welcome(self):
        u, _ = load_creds()
        if u:
            self._append(f"[INFO] מוכן. משתמש: {u}", "INFO")
            self._append("[INFO] הגדר ימים אחורה ולחץ ▶ הפעל", "INFO")
        else:
            self._append("[WARN] פרטי כניסה לא נשמרו — לחץ 🔑 פרטי כניסה", "WARN")
            self._append("[INFO] לאחר מכן הגדר ימים אחורה ולחץ ▶ הפעל", "INFO")

    def _build_ui(self):
        pad = {"padx": 10, "pady": 5}

        # Settings frame
        top = tk.LabelFrame(self, text="הגדרות", padx=8, pady=6)
        top.pack(fill="x", padx=10, pady=(10, 0))

        tk.Label(top, text="ימים אחורה:").grid(row=0, column=0, sticky="e", **pad)
        self._days = tk.StringVar(value="3")
        tk.Spinbox(top, from_=1, to=365, textvariable=self._days,
                   width=6, justify="center").grid(row=0, column=1, sticky="w", **pad)

        tk.Label(top, text="כמה כרטיסים לסרוק:").grid(row=1, column=0, sticky="e", **pad)
        self._max_docs = tk.StringVar(value="30")
        ttk.Combobox(top, textvariable=self._max_docs,
                     values=["10", "20", "30", "45", "60", "100"],
                     state="readonly", width=8).grid(row=1, column=1, sticky="w", **pad)

        # Mark-as-saved option
        self._mark_saved   = tk.BooleanVar(value=False)
        self._remember_mark = tk.BooleanVar(value=False)
        mark_frame = tk.Frame(top)
        mark_frame.grid(row=2, column=0, columnspan=3, sticky="w", padx=10, pady=(4, 0))
        tk.Checkbutton(
            mark_frame, text='?האם להחליף את הסטטוס ל"נשמר"',
            variable=self._mark_saved, command=self._on_mark_saved_toggle,
        ).pack(side="left")
        tk.Checkbutton(
            mark_frame, text="זכור העדפה",
            variable=self._remember_mark, command=self._on_remember_toggle,
            fg="#666666",
        ).pack(side="left", padx=(16, 0))

        # Buttons
        bf = tk.Frame(self)
        bf.pack(fill="x", padx=10, pady=8)

        self._run_btn = tk.Button(
            bf, text="▶  הפעל", width=14,
            bg="#2e7d32", fg="white", activebackground="#1b5e20",
            font=("Segoe UI", 10, "bold"), command=self._on_run,
        )
        self._run_btn.pack(side="left", padx=(0, 8))

        self._stop_btn = tk.Button(
            bf, text="■  עצור", width=10,
            bg="#c62828", fg="white", state="disabled",
            font=("Segoe UI", 10, "bold"), command=self._on_stop,
        )
        self._stop_btn.pack(side="left")

        tk.Button(bf, text="🔑  פרטי כניסה",
                  command=self._on_creds).pack(side="right")

        # Log
        lf = tk.LabelFrame(self, text="לוג בזמן אמת", padx=6, pady=4)
        lf.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self._log = scrolledtext.ScrolledText(
            lf, width=88, height=26, state="disabled",
            font=("Consolas", 9), bg="#1e1e1e", fg="#d4d4d4",
        )
        self._log.pack(fill="both", expand=True)
        self._log.tag_config("ERROR", foreground="#f44336")
        self._log.tag_config("WARN",  foreground="#ff9800")
        self._log.tag_config("INFO",  foreground="#d4d4d4")

        # Status bar
        self._status = tk.StringVar(value="מוכן")
        tk.Label(self, textvariable=self._status, anchor="w",
                 relief="sunken", bd=1).pack(fill="x", side="bottom")

    def _on_run(self):
        if not SELENIUM_OK:
            messagebox.showerror(
                "שגיאה",
                "Selenium לא מותקן.\n\npip install selenium webdriver-manager openpyxl keyring"
            )
            return
        try:
            days = int(self._days.get())
            assert days > 0
        except Exception:
            messagebox.showwarning("קלט שגוי", "נא להזין מספר ימים חיובי")
            return

        u, p = load_creds()
        if not u or not p:
            self._on_creds()
            u, p = load_creds()
            if not u or not p:
                return

        dlg = FolderDialog(self)
        if not dlg._confirmed:
            return

        self._clear_log()
        self._set_running(True)
        self._bot = Bot(
            days_back=days,
            invoice_folder=Path(dlg.invoice_folder),
            receipt_folder=Path(dlg.receipt_folder),
            log_cb=self._append,
            done_cb=self._on_done,
            max_docs=int(self._max_docs.get()),
            mark_as_saved=self._mark_saved.get(),
        )
        self._bot.start()

    def _on_stop(self):
        if self._bot:
            self._bot.stop()
        self._status.set("עוצר…")

    def _on_mark_saved_toggle(self):
        if self._remember_mark.get():
            self._persist_mark_pref()

    def _on_remember_toggle(self):
        self._persist_mark_pref()

    def _persist_mark_pref(self):
        cfg = load_folder_config()
        cfg['remember_mark_saved'] = self._remember_mark.get()
        cfg['mark_as_saved']       = self._mark_saved.get()
        save_folder_config(cfg)

    def _on_creds(self):
        dlg = CredsDialog(self)
        if dlg.username and dlg.password:
            save_creds(dlg.username, dlg.password)
            self._append("[INFO] פרטי כניסה נשמרו", "INFO")

    def _on_done(self, need_creds=False):
        self.after(0, self._set_running, False)
        if need_creds:
            self.after(0, self._on_creds)

    def _set_running(self, running: bool):
        self._run_btn.config(state="disabled" if running else "normal")
        self._stop_btn.config(state="normal"  if running else "disabled")
        self._status.set("פועל…" if running else "מוכן")

    def _clear_log(self):
        self._log.config(state="normal")
        self._log.delete("1.0", "end")
        self._log.config(state="disabled")

    def _append(self, message: str, level: str = "INFO"):
        def _do():
            self._log.config(state="normal")
            tag = level if level in ("ERROR", "WARN") else "INFO"
            self._log.insert("end", message + "\n", tag)
            self._log.see("end")
            self._log.config(state="disabled")
        self.after(0, _do)


if __name__ == "__main__":
    App().mainloop()
